import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import matplotlib.pyplot as plt
from io import BytesIO
import tkinter as tk
from tkinter import messagebox, simpledialog
from pathlib import Path
import sys
import os



def resource_path(relative_path):
    """ PyInstaller exe থেকে resource path ঠিক করতে """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
# Firebase Init
cred = credentials.Certificate(resource_path("google_services.json"))
firebase_admin.initialize_app(cred)
db = firestore.client()

# Bangla Options for MCQ options prefix
bangla_options = ['A', 'B', 'C', 'D', 'E', 'চ', 'ছ', 'জ']

# Latex rendering function (non-TeX mode)
def render_latex_image(latex_code):
    plt.rc('text', usetex=False)
    fig, ax = plt.subplots(figsize=(len(latex_code) * 0.25 + 1, 1))
    ax.text(0.5, 0.5, f"${latex_code}$", fontsize=20, ha='center', va='center')
    ax.axis('off')
    buf = BytesIO()
    plt.savefig(buf, format='PNG', bbox_inches='tight', pad_inches=0.1, transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf

def extract_latex(text):
    # Extract LaTeX between $$ and **
    match = re.search(r'\$\$(.*?)\*\*', text)
    if match:
        before = text[:match.start()]
        latex_code = match.group(1)
        after = text[match.end():]
        return before.strip(), latex_code.strip(), after.strip()
    return text, None, None

def add_text_with_latex(paragraph, raw_text, font_size=12):
    before, latex_code, after = extract_latex(raw_text)

    if before:
        run = paragraph.add_run(before + " ")
        run.font.name = 'NikoshBAN'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NikoshBAN')
        run.font.size = Pt(font_size)

    if latex_code:
        img_buf = render_latex_image(latex_code)
        run = paragraph.add_run()
        run.add_picture(img_buf, width=Inches(1.5))

    if after:
        run = paragraph.add_run(" " + after)
        run.font.name = 'NikoshBAN'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NikoshBAN')
        run.font.size = Pt(font_size)

def set_two_columns(doc):
    sectPr = doc.sections[0]._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set('num', '2')

def create_word_doc(mcqs, exam_name, filename="mcq_exam.docx"):
    doc = Document()
    set_two_columns(doc)
    doc.add_heading(f"MCQ প্রশ্নপত্র - {exam_name}", level=1)

    for idx, mcq in enumerate(mcqs, 1):
        q_para = doc.add_paragraph(f"{idx}. ")
        q_para.paragraph_format.space_after = Pt(1)
        q_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        add_text_with_latex(q_para, mcq.get("question", ""), font_size=14)

        options = mcq.get("options", [])
        option_para = doc.add_paragraph()
        option_para.paragraph_format.space_after = Pt(3)

        for i, opt in enumerate(options):
            prefix = f"{bangla_options[i]}. "
            run = option_para.add_run(prefix)
            run.font.name = 'NikoshBAN'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NikoshBAN')
            run.font.size = Pt(12)

            before, latex_code, after = extract_latex(opt)
            if before:
                run = option_para.add_run(before + " ")
                run.font.name = 'NikoshBAN'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NikoshBAN')
                run.font.size = Pt(12)
            if latex_code:
                img_buf = render_latex_image(latex_code)
                run = option_para.add_run()
                run.add_picture(img_buf, width=Inches(1.2))
            if after:
                run = option_para.add_run(" " + after + "   ")
                run.font.name = 'NikoshBAN'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'NikoshBAN')
                run.font.size = Pt(12)

        doc.add_paragraph()  # Blank paragraph for spacing

    # Save to Downloads folder
    downloads_path = str(Path.home() / "Downloads")
    filepath = f"{downloads_path}/{filename}"
    doc.save(filepath)
    messagebox.showinfo("✅ Success", f"Word file তৈরি হয়েছে:\n{filepath}")

def get_all_exams():
    exams_ref = db.collection("Exams").stream()
    return {doc.id: doc.to_dict().get("topicName", "Unnamed Exam") for doc in exams_ref}

def get_mcqs(exam_id, limit=None):
    ref = db.collection("Exams").document(exam_id).collection("MCQ")
    mcqs = [doc.to_dict() for doc in ref.stream()]
    return mcqs[:limit] if limit else mcqs

# --- GUI Start ---
def on_exam_select(event):
    selection = listbox.curselection()
    if selection:
        index = selection[0]
        exam_id = exam_ids[index]
        exam_name = exam_names[index]

        try:
            mcq_count = simpledialog.askinteger("MCQ Limit", f"'{exam_name}' থেকে কয়টা MCQ নিবে?", minvalue=1)
            if mcq_count:
                mcqs = get_mcqs(exam_id, mcq_count)
                create_word_doc(mcqs, exam_name, f"{exam_name}_MCQs.docx")
        except Exception as e:
            messagebox.showerror("❌ Error", str(e))

root = tk.Tk()
root.title("MCQ Exporter")
root.geometry("400x400")

tk.Label(root, text="Exam লিস্ট (Firestore থেকে)", font=("Helvetica", 14)).pack(pady=10)

listbox = tk.Listbox(root, font=("Helvetica", 12), height=15)
listbox.pack(fill=tk.BOTH, expand=True, padx=20)

exam_dict = get_all_exams()
exam_ids = list(exam_dict.keys())
exam_names = list(exam_dict.values())

for name in exam_names:
    listbox.insert(tk.END, name)

listbox.bind('<<ListboxSelect>>', on_exam_select)

root.mainloop()
